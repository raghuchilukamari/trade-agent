"""
Report Generator — creates professional DOCX trading analysis documents.

Uses python-docx for JP Morgan-standard formatting with color-coded tables,
ticker marks, and structured sections per the Trading Analysis Master System.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import structlog

from config.settings import settings

logger = structlog.get_logger(__name__)

# ── Color Constants (from SKILL.md) ──────────────────────────────────────────

COLORS = {
    "bullish": "166534",
    "bearish": "dc2626",
    "mixed": "ea580c",
    "info": "2563eb",
    "title": "1a365d",
    "heading1": "1e40af",
    "heading2": "374151",
    "gray": "666666",
    "purple": "7c3aed",
    "dark_gray": "4b5563",
    "black": "000000",
}

SHADING = {
    "bullish": "dcfce7",
    "bullish_light": "f0fdf4",
    "bearish": "fee2e2",
    "mixed": "fef3c7",
    "info": "dbeafe",
    "elevated": "ffedd5",
    "disclaimer": "f3f4f6",
}


async def generate_docx(report: dict[str, Any], context: dict[str, Any]) -> str:
    """
    Generate a DOCX report from the final analysis report.

    Args:
        report: FinalReport dict from coordinator
        context: PipelineContext dict

    Returns:
        Path to the generated DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx not installed")
        return ""

    doc = Document()
    target_date = context["target_date"]
    analysis_type = context["analysis_type"].upper()

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_heading("Daily Trading Update", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string(COLORS["title"])

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"{target_date} — {analysis_type} ANALYSIS")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(COLORS["heading1"])

    # ── Executive Summary ────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    exec_summary = report.get("executive_summary", "No summary available.")
    p = doc.add_paragraph()
    p.add_run(exec_summary).font.size = Pt(11)

    # ── OPEX Context Box ─────────────────────────────────────────────────
    opex = report.get("task4_opex", {})
    if opex:
        doc.add_heading("OPEX Context", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Next OPEX: {opex.get('next_monthly_opex', 'N/A')} "
                   f"({opex.get('days_to_opex', '?')} days)\n").bold = True
        p.add_run(f"Phase: {opex.get('phase_label', 'Unknown')}\n")
        p.add_run(f"Gamma: {opex.get('gamma_assessment', 'N/A')}\n")
        if opex.get("is_quad_witching"):
            run = p.add_run("⚠️ QUAD WITCHING\n")
            run.font.color.rgb = RGBColor.from_string(COLORS["bearish"])
            run.bold = True

    # ── Task 1: News-Flow Correlation ────────────────────────────────────
    verdicts = report.get("task1_correlation", [])
    if verdicts:
        doc.add_heading("Task 1: News-Flow Correlation", level=1)
        for v in verdicts:
            marks = "".join(v.get("marks", []))
            verdict_color = (
                COLORS["bullish"] if v.get("verdict") == "BULLISH"
                else COLORS["bearish"] if v.get("verdict") == "BEARISH"
                else COLORS["mixed"]
            )

            p = doc.add_paragraph()
            run = p.add_run(f"{v.get('symbol', '?')} {marks} — {v.get('verdict', '?')}")
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(verdict_color)
            run.font.size = Pt(12)

            p2 = doc.add_paragraph()
            p2.add_run("News: ").bold = True
            p2.add_run(f"{v.get('news_summary', '')} (Sentiment: {v.get('news_sentiment', 0):.1f}/5)\n")
            p2.add_run("Flow: ").bold = True
            p2.add_run(f"{v.get('flow_summary', '')}\n")
            p2.add_run("Verdict: ").bold = True
            p2.add_run(f"{v.get('reasoning', '')}")

    # ── Task 2: Geopolitical ─────────────────────────────────────────────
    geo = report.get("task2_geopolitical", {})
    if geo:
        doc.add_heading("Task 2: Geopolitical Analysis", level=1)

        bullish = geo.get("geopolitical_bullish", [])
        if bullish:
            doc.add_heading("Bullish Factors", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Topic"
            hdr[1].text = "Summary"
            hdr[2].text = "Score"
            for item in bullish:
                row = table.add_row().cells
                row[0].text = item.get("topic", "")
                row[1].text = item.get("summary", "")[:100]
                row[2].text = f"{item.get('score', 0):.1f}"

        bearish = geo.get("geopolitical_bearish", [])
        if bearish:
            doc.add_heading("Bearish Factors", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Topic"
            hdr[1].text = "Summary"
            hdr[2].text = "Score"
            for item in bearish:
                row = table.add_row().cells
                row[0].text = item.get("topic", "")
                row[1].text = item.get("summary", "")[:100]
                row[2].text = f"{item.get('score', 0):.1f}"

    # ── Task 3: Top 10 Flow Trades ───────────────────────────────────────
    top_trades = report.get("task3_top_trades", [])
    if top_trades:
        doc.add_heading("Task 3: Top 10 Flow Trades", level=1)
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["#", "Symbol", "Type", "Strike", "Exp", "Premium", "Vol/OI", "Notes"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for idx, t in enumerate(top_trades[:10], 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = f"{t.get('symbol', '')} {t.get('marks_str', '')}"
            row[2].text = t.get("call_put", "")
            row[3].text = f"${t.get('strike', '')}" if t.get("strike") else ""
            row[4].text = str(t.get("expiration", ""))
            row[5].text = t.get("premium_formatted", "")
            row[6].text = f"{t.get('vol_oi', '')}" if t.get("vol_oi") else ""
            notes = []
            if t.get("deep_itm_note"):
                notes.append(t["deep_itm_note"])
            if t.get("news_correlated"):
                notes.append("📰 News correlated")
            row[7].text = "; ".join(notes) if notes else t.get("source", "")

    # ── Risk Assessment ──────────────────────────────────────────────────
    risks = report.get("risk_assessment", [])
    if risks:
        doc.add_heading("Risk Assessment", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, h in enumerate(["Category", "Level", "Description"]):
            table.rows[0].cells[i].text = h
        for r in risks:
            row = table.add_row().cells
            row[0].text = r.get("category", "")
            row[1].text = r.get("level", "")
            row[2].text = r.get("description", "")

    # ── Action Items ─────────────────────────────────────────────────────
    actions = report.get("action_items", [])
    if actions:
        doc.add_heading("Action Items", level=1)
        for a in actions:
            doc.add_paragraph(a, style="List Bullet")

    # ── Sources ──────────────────────────────────────────────────────────
    doc.add_heading("Sources & Citations", level=1)
    for src in report.get("sources", []):
        doc.add_paragraph(src, style="List Bullet")

    # ── Disclaimer ───────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "DISCLAIMER: This analysis is for informational purposes only and does not "
        "constitute financial advice. Options trading involves significant risk. "
        "Past performance is not indicative of future results."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS["gray"])

    # ── Save ─────────────────────────────────────────────────────────────
    output_dir = Path(settings.summary_output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"trading_update_{target_date}_{context['analysis_type']}.docx"
    filepath = output_dir / filename

    doc.save(str(filepath))
    logger.info("docx_saved", path=str(filepath))

    return str(filepath)
